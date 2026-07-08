"""Word (.docx) parser.

Extracts text from every non-empty paragraph (including those inside tables)
and rewrites them in place, preserving the document's overall structure. Run
formatting is simplified: the translated text is placed into the paragraph's
first run and remaining runs are cleared, which keeps paragraph-level style.
"""

from __future__ import annotations

from .base import DocumentParser, register


@register
class DocxParser(DocumentParser):
    extensions = (".docx",)

    def __init__(self, path):
        super().__init__(path)
        self._doc = None
        self._paragraphs = []  # translatable paragraph objects in order

    def _iter_paragraphs(self, doc):
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    def extract(self) -> list[str]:
        from docx import Document

        self._doc = Document(str(self.path))
        self._paragraphs = [
            p for p in self._iter_paragraphs(self._doc) if p.text.strip()
        ]
        return [p.text for p in self._paragraphs]

    def rebuild(self, translations: list[str], output_path) -> None:
        for para, translated in zip(self._paragraphs, translations):
            runs = para.runs
            if runs:
                runs[0].text = translated
                for run in runs[1:]:
                    run.text = ""
            else:
                para.add_run(translated)
        self._doc.save(str(output_path))
